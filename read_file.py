import pandas as pd
import numpy as np
from metric_converter import Metric
import docx as _readFile
from docx2python import docx2python
# from docx2python.docx_context import collect_numFmts
# from metric_converter import Converter
from docx.enum.style import WD_STYLE_TYPE


def file_base_style(filename):
    doc = _readFile.Document(filename)
    styles = doc.styles
    paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    for style in paragraph_styles:
        print(style.name)


def get_resources(filename):
    doc = _readFile.Document(filename)
    paragraphs = doc.paragraphs
    for paragraph in paragraphs:
        if(paragraph.text == "KAYNAKLAR"):
            print("----------------------KAYNAKLAR----------------------")
        if(paragraph.text != "" and paragraph.style.name == 'List Paragraph'):
            print(paragraph.style.name, "---", paragraph.text, "\n")

    # def file_body(filename):
    #    doc = docx2python(filename)
    #    return doc.body

    #    #print('body', doc.body[0][0][0][1])


""" DOCX İLE YAZDIKLARIM"""
# sections = doc.sections
# for i in sections:
#    print(i.start_type)

# buradaki return edilen dict i html de tablo oluşturarak ekrana yaz
# tablodaki her bir değeri şu şekilde alabilirsin
# for row in table.rows:
#    for cell in row.cells:
#        print(cell.text)
# [[cell.text for cell in row.cells] for row in table.rows]


# Bu dınıfın isimlendirme kuralı: def lerin ilk kelimesi read_ ile başlayacak
class WithDocx:
    """
    ENGLISH:\n
        There are methods I wrote with the docx library.\n
    TURKÇE:\n
        Docx kütüphanesi ile yazdığım metotlar bulunur.\n
    """
    def read_table(filename):
        """
        ENGLISH:\n
            It provides access to the tables in the file. And it enables us to get a meaningful result by using pandas and numpy libraries.
        TÜRKÇE:\n
            Dosyada bulunan tablolara ulaşmayı sağlar. Ve pandas ile numpy kütüphanelerinden faydalanarak anlamlı bir sonuç elde etmemizi sağlar.
        Args:\n
            filename (text):\n
                ENGLISH:\n
                    File name.
                TÜRKÇE:\n
                    Dosya adı.
        Returns:\n
            list :\n
                ENGLISH:\n
                    Returns a nested list containing table data. If there is no table in the file, it returns None.
                TÜRKÇE:\n
                    Tablo verisini içeren iç içe liste döndürür. Eğer dosyada tablo yoksa None döndürür.
        """
        doc = _readFile.Document(filename)
        tables = doc.tables
        toTable = []
        # bulunan tablo sayısı. len(tables)
        for table in tables:
            toTable.append(pd.DataFrame(np.array([[table.rows[row].cells[column].text for column in range(
                len(table.columns))] for row in range(len(table.rows))])))
        if toTable == []:
            return None
        else:
            return toTable

    def read_paragraph(filename):
        """
        ENGLISH:
            Returns each paragraph in the file as a sequence.
        TÜRKÇE:
            Dosyadaki her bir paragrafı bir dizi olarak geri döndürür.
        \nArgs:
            filename (str):
                ENGLISH:
                    File path.
                TÜRKÇE:
                    Dosya yolu.
        \nReturns:
            str:
            ENGLISH:
                Returns each line in the given file as a string.
            TÜRKÇE:
                Verilen dosyadaki her bir satırı dizi şeklinde geri döndürür.
        """
        doc = _readFile.Document(filename)
        return [p.text for p in doc.paragraphs]

    # None left oluyor standart olan
    # eğer None ve i.text i null ise enter tuşu oluyor alt satıra geçmek
    # CENTER (1)
    # RIGHT (2)
    # JUSTIFY (3)

    def read_alignment(filename):
        """
        ENGLISH:
            It is used to find the alignment values of each paragraph in the file.
        TÜRKÇE:
            Dosyadaki her bir paragrafın hizalama değerlerini bulmaya yarar.
        \nArgs:
            filename (str):
                ENGLISH:
                    File path.
                TÜRKÇE:
                    Dosya yolu.
        \nReturns:
            str:
            ENGLISH:
                Returns the alignment values as a dictionary data structure.
            TÜRKÇE:
                Hizalama değerlerini sözlük veri yapısı şeklinde geri döndürür.
            `mainAlignmentArray['alignment']`
            `mainAlignmentArray['text']`
        """
        mainAlignmentArray = {}
        paragraphAlignmentArray = []
        paragraphTextArray = []
        doc = _readFile.Document(filename)
        paragraph = doc.paragraphs
        for i in paragraph:
            paragraphAlignmentArray.append(i.paragraph_format.alignment)
            paragraphTextArray.append(i.text)
        mainAlignmentArray['alignment'] = paragraphAlignmentArray
        mainAlignmentArray['text'] = paragraphTextArray
        return mainAlignmentArray

    def read_page_structure(filename):
        """
        ENGLISH:
            It is used to find the margin values of each paragraph in the file.
        TÜRKÇE:
            Dosyadaki her bir paragrafın kenar boşluk değerlerini bulmaya yarar.
        \nArgs:
            filename (str):
                ENGLISH:
                    File path.
                TÜRKÇE:
                    Dosya yolu.
        \nReturns:
            str:
            ENGLISH:
                Returns the margin values as a dictionary data structure.
            TÜRKÇE:
                Kenar boşluk değerlerini sözlük veri yapısı şeklinde geri döndürür.
            `mainMarginArray['top']`
            `mainMarginArray['right']`
            `mainMarginArray['bottom']`
            `mainMarginArray['left']`
            `mainMarginArray['portrait']`
        """
        mainMarginArray = {}
        topMarginArray = []
        rightMarginArray = []
        bottomMarginArray = []
        leftMarginArray = []
        portraitMarginArray = []
        doc = _readFile.Document(filename)
        sections = doc.sections
        for i in sections:
            if(str(i.start_type) == 'NEW_PAGE (2)'):
                if(str(i.orientation) == 'PORTRAIT (0)'):
                    portraitMarginArray.append('Dikey')
                if(str(i.orientation) == 'LANDSCAPE (1)'):
                    portraitMarginArray.append('Yatay')
                topMarginArray.append(round(Metric.emuToCm(i.top_margin), 2))
                rightMarginArray.append(
                    round(Metric.emuToCm(i.right_margin), 2))
                bottomMarginArray.append(
                    round(Metric.emuToCm(i.bottom_margin), 2))
                leftMarginArray.append(round(Metric.emuToCm(i.left_margin), 2))
        mainMarginArray['top'] = topMarginArray
        mainMarginArray['right'] = rightMarginArray
        mainMarginArray['bottom'] = bottomMarginArray
        mainMarginArray['left'] = leftMarginArray
        mainMarginArray['portrait'] = portraitMarginArray
        return mainMarginArray

    def read_alignment_and_indent(filename):
        """
        ENGLISH:
            It is used to find the indent values of each paragraph in the file.
        TÜRKÇE:
            Dosyadaki her bir paragrafın girinti değerlerini bulmaya yarar.
        \nArgs:
            filename (str):
                ENGLISH:
                    File path.
                TÜRKÇE:
                    Dosya yolu.
        \nReturns:
            str:
            ENGLISH:
                Returns the indent values as a dictionary data structure.
            TÜRKÇE:
                Girinti değerlerini sözlük veri yapısı şeklinde geri döndürür.
            `read_alignment_and_indent['left']`
            `read_alignment_and_indent['right']`
        """
        mainIndentArray = {}
        leftIndent = []
        rightIndent = []
        doc = _readFile.Document(filename)
        indent = doc.paragraphs
        for i in indent:
            leftIndent.append(round(Metric.emuToCm(
                i.paragraph_format.left_indent), 2))
            rightIndent.append(round(Metric.emuToCm(
                i.paragraph_format.right_indent), 2))
        mainIndentArray['left'] = leftIndent
        mainIndentArray['right'] = rightIndent
        return mainIndentArray


"""DOCX2PYTHON İLE YAZDIKLARIM"""


# Bu dınıfın isimlendirme kuralı: def lerin ilk kelimesi read_ ile başlayacak

class WithDocx2Python:

    def file_save_image(filename, image_folder):
        doc = docx2python(filename, image_folder)
        return doc.properties

    def file_properties(filename):
        """
        ENGLISH:
            It is a method that provides access to file properties.
        TÜRKÇE:
            Dosya özelliklerine ulaşmayı sağlayan metoddur.
        \nArgs:
            filename (str):
                ENGLISH:
                    File path.
                TÜRKÇE:
                    Dosya yolu.
        \nReturns:
            str:
            ENGLISH:
                Returns a dictionary data structure containing file properties.
            TÜRKÇE:
                Dosya özelliklerini içeren bir sözlük veri yapısı döndürür.
        """
        doc = docx2python(filename)
        return doc.properties

    def file_header(filename):
        """
        ENGLISH:
            It is the method that provides access to the header of the file.
        TÜRKÇE:
            Dosyanın üst başlığını ulaşmayı sağlayan metoddur.
        \nArgs:
            filename (str):
                ENGLISH:
                    File path.
                TÜRKÇE:
                    Dosya yolu.
        \nReturns:
            str:
            ENGLISH:
                Returns a TablesList containing file properties.
            TÜRKÇE:
                Dosya özelliklerini içeren bir TablesList döndürür.
        """
        doc = docx2python(filename)
        return doc.header

    def file_footer(filename):
        """
        ENGLISH:
            It is the method that provides access to the footer of the file.
        TÜRKÇE:
            Dosyanın alt başlığını ulaşmayı sağlayan metoddur.
        \nArgs:
            filename (str):
                ENGLISH:
                    File path.
                TÜRKÇE:
                    Dosya yolu.
        \nReturns:
            str:
            ENGLISH:
                Returns a TablesList containing file properties.
            TÜRKÇE:
                Dosya özelliklerini içeren bir TablesList döndürür.
        """
        doc = docx2python(filename)
        return doc.footer

    # Dönüş türü [] list
    # genelde bütün data bu kısımda oluyor
